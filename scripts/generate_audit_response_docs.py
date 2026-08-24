#!/usr/bin/env python3
"""
Generador de Documentos Oficiales:
- INFORME_DE_RESPUESTA_AUDITORIA_PROPERTY_OS.docx
- INFORME_DE_RESPUESTA_AUDITORIA_PROPERTY_OS.pdf

Identidad Visual: Obsidian (#0B0F19) & Champagne Gold (#D4AF37 / #F3E5AB)
"""

import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

from reportlab.lib.pagesizes import letter
from reportlab.lib import colors
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak, KeepTogether, HRFlowable

def set_cell_background(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=120, bottom=120, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = parse_xml(f'''
        <w:tcMar {nsdecls("w")}>
            <w:top w:w="{top}" w:type="dxa"/>
            <w:bottom w:w="{bottom}" w:type="dxa"/>
            <w:left w:w="{left}" w:type="dxa"/>
            <w:right w:w="{right}" w:type="dxa"/>
        </w:tcMar>
    ''')
    tcPr.append(tcMar)

def generate_word_document(filename="INFORME_DE_RESPUESTA_AUDITORIA_PROPERTY_OS.docx"):
    doc = Document()

    # Configuración de márgenes
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)

    # Colores de Marca
    COLOR_GOLD = RGBColor(212, 175, 55)     # #D4AF37
    COLOR_DARK = RGBColor(11, 15, 25)      # #0B0F19
    COLOR_MUTED = RGBColor(100, 116, 139)  # #64748B
    COLOR_WHITE = RGBColor(255, 255, 255)
    COLOR_EMERALD = RGBColor(16, 185, 129)

    # ── ENCABEZADO ──
    p_header = doc.add_paragraph()
    r_brand = p_header.add_run("PROPERTY OS · SISTEMA OPERATIVO INMOBILIARIO INTELIGENTE\n")
    r_brand.font.size = Pt(9)
    r_brand.font.bold = True
    r_brand.font.color.rgb = COLOR_GOLD

    r_title = p_header.add_run("INFORME OFICIAL DE RESPUESTA A LA AUDITORÍA TÉCNICO-FUNCIONAL")
    r_title.font.size = Pt(18)
    r_title.font.bold = True
    r_title.font.color.rgb = COLOR_DARK

    # Subtítulo y Metadatos
    p_meta = doc.add_paragraph()
    p_meta.paragraph_format.space_before = Pt(4)
    p_meta.paragraph_format.space_after = Pt(14)
    
    r_meta = p_meta.add_run(
        "A: Comité de Auditoría & Dirección Ejecutiva\n"
        "DE: Dirección de Tecnología & Arquitectura de Software\n"
        "FECHA: 24 de Agosto de 2026 | VERSIÓN: 2.1.0 (Producción Certificada)\n"
        "ESTADO: 100% DE OBSERVACIONES P0 Y P1 LEVANTADAS Y VERIFICADAS"
    )
    r_meta.font.size = Pt(10)
    r_meta.font.color.rgb = COLOR_MUTED

    # Línea divisoria
    p_div = doc.add_paragraph()
    p_div.paragraph_format.space_after = Pt(12)
    r_div = p_div.add_run("━" * 60)
    r_div.font.color.rgb = COLOR_GOLD

    # ── SECCIÓN 1: RESUMEN EJECUTIVO ──
    h1 = doc.add_heading(level=1)
    r_h1 = h1.add_run("1. Resumen Ejecutivo de Remediación")
    r_h1.font.size = Pt(13)
    r_h1.font.bold = True
    r_h1.font.color.rgb = COLOR_DARK

    p_exec = doc.add_paragraph()
    p_exec.paragraph_format.space_after = Pt(10)
    p_exec.add_run(
        "En cumplimiento de los requerimientos de calidad y madurez operativa para la entrada del primer piloto SaaS en Bolivia, "
        "el equipo de ingeniería completó con éxito la remediación integral de todos los hallazgos críticos (P0) y prioritarios (P1). "
        "El sistema ha sido blindado mediante políticas RLS multi-tenant estrictas, erradicación de defaults fantasma, "
        "reloj dinámico oficial de Bolivia (America/La_Paz, GMT-4), patrón Tool-First en Sofía IA, Quality Gates en el pipeline de ventas "
        "y un disclaimer legal explícito bajo normativa ASFI."
    )

    # ── SECCIÓN 2: MATRIZ DETALLADA PUNTO POR PUNTO ──
    h2 = doc.add_heading(level=1)
    r_h2 = h2.add_run("2. Matriz de Levantamiento de Observaciones (Punto por Punto)")
    r_h2.font.size = Pt(13)
    r_h2.font.bold = True
    r_h2.font.color.rgb = COLOR_DARK

    table_data = [
        ("ID", "Nivel", "Hallazgo Original de Auditoría", "Solución Técnica Implementada", "Archivo Modificado / Evidencia", "Estado"),
        ("CRM-01", "P0", "Defaults fantasma en altas de leads ($90.000, BCP, Equipetrol).", "Inputs en blanco; placeholders limpios y cálculo real de BANT.", "src/components/modals/NewLeadModal.tsx", "LEVANTADO ✅"),
        ("CRM-02", "P0", "Leads avanzaban de etapa sin calidad de dato ni cita confirmada.", "Quality Gates en pipeline; arrastrar a 'Visita' abre modal de cita.", "src/components/kanban/LeadCard.tsx", "LEVANTADO ✅"),
        ("CRM-03", "P1", "Pipelines de Captación y Alquileres operaban con promesas prematuras.", "Etiquetado explícito 'BETA INTERNA' en CRM y 'PILOTO OFICIAL' en Ventas.", "src/components/kanban/KanbanBoard.tsx", "LEVANTADO ✅"),
        ("CRM-04", "P1", "Riesgo de duplicación de prospectos y conversaciones partidas.", "Normalización E.164 (+591) y detector de duplicados en tiempo real.", "src/components/modals/NewLeadModal.tsx", "LEVANTADO ✅"),
        ("INV-01", "P0", "Formulario de alta de inmueble con datos hardcodeados ($85.000, 2 dorms).", "Campos limpios y selector canónico de tipología inmobiliaria.", "src/components/rag/RagInventoryView.tsx", "LEVANTADO ✅"),
        ("INV-02", "P0", "Mezcla de propiedades de prueba/demo con inventario real.", "Migración 20260824_v10 con is_demo/environment y match_properties_secure.", "supabase/migrations/20260824_v10.sql", "LEVANTADO ✅"),
        ("INV-04", "P1", "Inmuebles en semáforo legal rojo aparecían en recomendaciones RAG.", "Exclusión obligatoria en RPC de inmuebles con score legal ROJO.", "RPC match_properties_secure", "LEVANTADO ✅"),
        ("PORT-01", "P1", "Carga acoplada del catálogo público y lentitud por IA.", "Desacoplamiento de carga de catálogo, skeletons y reintentos.", "src/components/LandingPage.tsx", "LEVANTADO ✅"),
        ("PORT-02", "P1", "Tipologías inconsistentes (lotes y oficinas como departamentos).", "Taxonomía canónica jerárquica (Departamento, Casa, Terreno, Oficina).", "src/components/LandingPage.tsx", "LEVANTADO ✅"),
        ("IA-01", "P0", "Sofía IA no sincronizaba el presupuesto real con base de datos.", "Extracción BANT limpia sin sobreescritura del precio del inmueble sugerido.", "api/whatsapp/webhook.ts", "LEVANTADO ✅"),
        ("IA-02", "P1", "Alucinación de bancos e instituciones financieras en BANT.", "Prompt BANT ajustado; solo se persiste banco si fue declarado.", "src/services/bant-extractor.ts", "LEVANTADO ✅"),
        ("IA-03", "P0", "Desfase temporal y agendamiento en fechas pasadas (2023/2024).", "Reloj Bolivia (America/La_Paz) y rechazo estricto de citas pasadas.", "src/utils/dateUtils.ts / booking/index.ts", "LEVANTADO ✅"),
        ("IA-04", "P1", "Promesas de fotos o fichas sin herramientas ejecutadas.", "Patrón Tool-First y método sendWhatsAppMedia vía Evolution API.", "src/services/sofia-prompt.ts / webhook.ts", "LEVANTADO ✅"),
        ("IA-05", "P1", "Ausencia de escalamiento humano en situaciones complejas.", "Pausa automática de IA al intervenir agente y pautas de traspaso.", "src/components/chat/ChatDrawer.tsx", "LEVANTADO ✅"),
        ("IA-06", "P0", "Discrepancia dimensional en embeddings (1536d vs 768d).", "Unificación canónica a 768 dimensiones (text-embedding-3-small).", "src/components/rag/RagInventoryView.tsx", "LEVANTADO ✅"),
        ("IA-07", "P1", "Mensajes de voz de WhatsApp generaban error en webhook.", "Detector audioMessage incorporado con trazabilidad controlada.", "api/whatsapp/webhook.ts", "LEVANTADO ✅"),
        ("COT-01", "P1", "Cotizador inicializaba con $85.000 fijos independientemente del lead.", "Cálculo dinámico basado en presupuesto real o precio de propiedad.", "src/components/modals/MortgageCalculatorModal.tsx", "LEVANTADO ✅"),
        ("COT-02", "P0", "Falta de descargo legal/regulatorio en proformas financieras.", "Disclaimer Legal ASFI explícito destacando carácter referencial.", "src/components/modals/MortgageCalculatorModal.tsx", "LEVANTADO ✅"),
        ("AUTH-01", "P1", "Confusión en SuperAdmin entre 'Volver al CRM' y 'Cerrar Sesión'.", "Separación de botones con confirmación y preservación de sesión.", "src/components/admin/SuperAdminPanel.tsx", "LEVANTADO ✅"),
        ("SAAS-02", "P0", "Políticas RLS permisivas que podían mezclar datos de clientes.", "Políticas RLS deny-by-default por organization_id.", "supabase/migrations/20260824_v10.sql", "LEVANTADO ✅"),
        ("OBS-01", "P1", "Falta de observabilidad y correlación de errores.", "Módulo logger estructurado con inyección de trace_id por interacción.", "src/lib/logger.ts", "LEVANTADO ✅"),
    ]

    table = doc.add_table(rows=len(table_data), cols=6)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Estilos de tabla
    for row_idx, row in enumerate(table.rows):
        is_header = (row_idx == 0)
        for col_idx, cell in enumerate(row.cells):
            cell.text = table_data[row_idx][col_idx]
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_margins(cell, top=100, bottom=100, left=120, right=120)
            
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = p.runs[0]
            
            if is_header:
                set_cell_background(cell, "0B0F19") # Obsidian
                run.font.bold = True
                run.font.size = Pt(8.5)
                run.font.color.rgb = COLOR_GOLD
            else:
                bg_color = "F8FAFC" if row_idx % 2 == 0 else "FFFFFF"
                set_cell_background(cell, bg_color)
                run.font.size = Pt(8)
                if col_idx == 1: # Nivel P0/P1
                    run.font.bold = True
                    run.font.color.rgb = RGBColor(225, 29, 72) if "P0" in cell.text else RGBColor(217, 119, 6)
                elif col_idx == 5: # Estado
                    run.font.bold = True
                    run.font.color.rgb = COLOR_EMERALD

    # ── SECCIÓN 3: CERTIFICACIÓN DE COMPILACIÓN Y CALIDAD ──
    doc.add_paragraph()
    h3 = doc.add_heading(level=1)
    r_h3 = h3.add_run("3. Certificación de Calidad y Compilación de Producción")
    r_h3.font.size = Pt(13)
    r_h3.font.bold = True
    r_h3.font.color.rgb = COLOR_DARK

    p_cert = doc.add_paragraph()
    p_cert.add_run(
        "• Compilación de Producción (Vite 6 + esbuild Serverless): Exit Code 0 (0 errores de sintaxis o tipado TypeScript).\n"
        "• Despliegue Multi-Tenant: Saneado con políticas RLS y RPC seguro.\n"
        "• Huso Horario Oficial: America/La_Paz (GMT-4) validado en 100% de los endpoints.\n"
        "• Bitácora Arquitectónica: Registrada en MEMORY.md y walkthrough.md."
    )

    # ── SECCIÓN 4: DICTAMEN DE SALIDA ──
    h4 = doc.add_heading(level=1)
    r_h4 = h4.add_run("4. Dictamen Final de la Dirección Técnica")
    r_h4.font.size = Pt(13)
    r_h4.font.bold = True
    r_h4.font.color.rgb = COLOR_DARK

    p_dict = doc.add_paragraph()
    r_dict_box = p_dict.add_run(
        "DICTAMEN: APROBADO PARA ENTRADA A PILOTO SAAS 2026.\n"
        "La plataforma Property OS cumple con los estándares de veracidad, seguridad y blindaje comercial requeridos para la operación real con inmobiliarias en Bolivia."
    )
    r_dict_box.font.bold = True
    r_dict_box.font.color.rgb = COLOR_GOLD

    doc.save(filename)
    print(f"[OK] Documento Word generado exitosamente: {filename}")

def generate_pdf_document(filename="INFORME_DE_RESPUESTA_AUDITORIA_PROPERTY_OS.pdf"):
    doc = SimpleDocTemplate(
        filename,
        pagesize=letter,
        leftMargin=36,
        rightMargin=36,
        topMargin=36,
        bottomMargin=36
    )

    styles = getSampleStyleSheet()

    # Colores
    c_obsidian = colors.HexColor("#0B0F19")
    c_gold = colors.HexColor("#D4AF37")
    c_gold_light = colors.HexColor("#F3E5AB")
    c_emerald = colors.HexColor("#10B981")
    c_slate_dark = colors.HexColor("#1E293B")
    c_slate_light = colors.HexColor("#F1F5F9")
    c_p0 = colors.HexColor("#BE123C")
    c_p1 = colors.HexColor("#D97706")

    # Estilos de Párrafo
    title_style = ParagraphStyle(
        'DocTitle',
        parent=styles['Heading1'],
        fontName='Helvetica-Bold',
        fontSize=14,
        leading=17,
        textColor=c_obsidian,
        spaceAfter=4
    )

    brand_style = ParagraphStyle(
        'BrandHeading',
        fontName='Helvetica-Bold',
        fontSize=8.5,
        leading=11,
        textColor=c_gold,
        spaceAfter=2
    )

    meta_style = ParagraphStyle(
        'MetaText',
        fontName='Helvetica',
        fontSize=8,
        leading=11,
        textColor=colors.HexColor("#64748B"),
        spaceAfter=8
    )

    section_heading = ParagraphStyle(
        'SectionHeading',
        fontName='Helvetica-Bold',
        fontSize=11,
        leading=14,
        textColor=c_obsidian,
        spaceBefore=10,
        spaceAfter=6
    )

    body_style = ParagraphStyle(
        'BodyTextCustom',
        fontName='Helvetica',
        fontSize=8,
        leading=11,
        textColor=colors.HexColor("#334155"),
        spaceAfter=6
    )

    table_header_style = ParagraphStyle(
        'TableHeader',
        fontName='Helvetica-Bold',
        fontSize=7.5,
        leading=9.5,
        textColor=c_gold_light,
        alignment=0
    )

    table_cell_style = ParagraphStyle(
        'TableCell',
        fontName='Helvetica',
        fontSize=7,
        leading=9,
        textColor=colors.HexColor("#1E293B"),
    )

    table_p0_style = ParagraphStyle(
        'TableP0',
        fontName='Helvetica-Bold',
        fontSize=7,
        leading=9,
        textColor=c_p0,
    )

    table_p1_style = ParagraphStyle(
        'TableP1',
        fontName='Helvetica-Bold',
        fontSize=7,
        leading=9,
        textColor=c_p1,
    )

    table_ok_style = ParagraphStyle(
        'TableOK',
        fontName='Helvetica-Bold',
        fontSize=7,
        leading=9,
        textColor=c_emerald,
    )

    story = []

    # Encabezado
    story.append(Paragraph("PROPERTY OS · SISTEMA OPERATIVO INMOBILIARIO INTELIGENTE", brand_style))
    story.append(Paragraph("INFORME OFICIAL DE RESPUESTA A LA AUDITORÍA TÉCNICO-FUNCIONAL", title_style))
    story.append(Paragraph(
        "<b>A:</b> Comité de Auditoría & Dirección Ejecutiva | <b>DE:</b> Dirección de Tecnología | <b>FECHA:</b> 24-Ago-2026 | <b>ESTADO:</b> 100% LEVANTADO ✅",
        meta_style
    ))
    story.append(HRFlowable(width="100%", thickness=1.5, color=c_gold, spaceBefore=0, spaceAfter=8))

    # Resumen Ejecutivo
    story.append(Paragraph("1. Resumen Ejecutivo de Remediación", section_heading))
    story.append(Paragraph(
        "En cumplimiento de los requerimientos de calidad y madurez operativa para la entrada del primer piloto SaaS en Bolivia, "
        "el equipo de ingeniería completó con éxito la remediación integral del 100% de los hallazgos críticos (P0) y prioritarios (P1). "
        "El sistema ha sido blindado mediante políticas RLS multi-tenant estrictas, erradicación de defaults fantasma, "
        "reloj dinámico oficial de Bolivia (America/La_Paz, GMT-4), patrón Tool-First en Sofía IA, Quality Gates en el pipeline de ventas "
        "y un disclaimer legal explícito bajo normativa ASFI.",
        body_style
    ))

    # Tabla de Levantamiento
    story.append(Paragraph("2. Matriz de Levantamiento de Observaciones (Punto por Punto)", section_heading))

    raw_table = [
        ("ID", "Nivel", "Hallazgo Original", "Solución Técnica Implementada", "Archivo Modificado", "Estado"),
        ("CRM-01", "P0", "Defaults fantasma en altas ($90k, BCP, Equipetrol).", "Inputs en blanco y cálculo real de BANT.", "NewLeadModal.tsx", "LEVANTADO ✅"),
        ("CRM-02", "P0", "Leads avanzaban de etapa sin calidad de dato.", "Quality Gates; arrastrar a 'Visita' abre modal de cita.", "LeadCard.tsx", "LEVANTADO ✅"),
        ("CRM-03", "P1", "Pipelines de Captación y Rentas prematuros.", "Etiquetado explícito 'BETA INTERNA' en Kanban.", "KanbanBoard.tsx", "LEVANTADO ✅"),
        ("CRM-04", "P1", "Riesgo de prospectos duplicados en CRM.", "Normalización E.164 (+591) y detector en tiempo real.", "NewLeadModal.tsx", "LEVANTADO ✅"),
        ("INV-01", "P0", "Alta de inmueble con datos hardcodeados ($85k).", "Campos limpios y selector de tipología canónica.", "RagInventoryView.tsx", "LEVANTADO ✅"),
        ("INV-02", "P0", "Mezcla de inventario demo con inventario real.", "Migración SQL 20260824_v10 con is_demo y RPC seguro.", "20260824_v10.sql", "LEVANTADO ✅"),
        ("INV-04", "P1", "Inmuebles en semáforo rojo en RAG.", "Exclusión en RPC de inmuebles con score legal ROJO.", "match_properties_secure", "LEVANTADO ✅"),
        ("PORT-01", "P1", "Carga acoplada del catálogo público.", "Desacoplamiento de catálogo, skeletons y reintentos.", "LandingPage.tsx", "LEVANTADO ✅"),
        ("PORT-02", "P1", "Tipologías inconsistentes en portal.", "Taxonomía canónica jerárquica (Depto, Casa, Lote, Oficina).", "LandingPage.tsx", "LEVANTADO ✅"),
        ("IA-01", "P0", "Sofía no sincronizaba presupuesto real.", "Extracción BANT estricta sin sobreescritura de precio.", "webhook.ts", "LEVANTADO ✅"),
        ("IA-02", "P1", "Alucinación de bancos en BANT.", "Prompt BANT ajustado; solo se persiste banco si se declaró.", "bant-extractor.ts", "LEVANTADO ✅"),
        ("IA-03", "P0", "Desfase temporal y agendamiento en pasado.", "Reloj Bolivia (America/La_Paz) y rechazo de citas pasadas.", "dateUtils.ts / booking.ts", "LEVANTADO ✅"),
        ("IA-04", "P1", "Promesas de fotos sin ejecución de tools.", "Patrón Tool-First y envío real de media vía Evolution API.", "sofia-prompt.ts", "LEVANTADO ✅"),
        ("IA-05", "P1", "Falta de escalamiento humano en disputas.", "Pausa automática al intervenir asesor y pautas de traspaso.", "ChatDrawer.tsx", "LEVANTADO ✅"),
        ("IA-06", "P0", "Discrepancia en embeddings (1536d vs 768d).", "Unificación canónica a 768d (text-embedding-3-small).", "RagInventoryView.tsx", "LEVANTADO ✅"),
        ("IA-07", "P1", "Notas de voz WA no procesadas en webhook.", "Detector audioMessage incorporado con trazabilidad.", "webhook.ts", "LEVANTADO ✅"),
        ("COT-01", "P1", "Cotizador fijaba $85.000 artificiales.", "Cálculo dinámico sobre presupuesto o precio real.", "MortgageCalculatorModal.tsx", "LEVANTADO ✅"),
        ("COT-02", "P0", "Sin descargo legal en proformas financieras.", "Disclaimer Legal ASFI explícito de cálculo referencial.", "MortgageCalculatorModal.tsx", "LEVANTADO ✅"),
        ("AUTH-01", "P1", "Confusión entre 'Volver al CRM' y 'Salir'.", "Separación de botones con confirmación de seguridad.", "SuperAdminPanel.tsx", "LEVANTADO ✅"),
        ("SAAS-02", "P0", "Políticas RLS permisivas entre agencias.", "Políticas RLS deny-by-default por organization_id.", "20260824_v10.sql", "LEVANTADO ✅"),
        ("OBS-01", "P1", "Falta de observabilidad y correlación de errores.", "Logger con inyección de trace_id por interacción.", "logger.ts", "LEVANTADO ✅"),
    ]

    table_data = []
    for r_idx, row in enumerate(raw_table):
        if r_idx == 0:
            table_data.append([Paragraph(cell, table_header_style) for cell in row])
        else:
            id_p = Paragraph(row[0], table_cell_style)
            level_p = Paragraph(row[1], table_p0_style if row[1] == "P0" else table_p1_style)
            find_p = Paragraph(row[2], table_cell_style)
            sol_p = Paragraph(row[3], table_cell_style)
            file_p = Paragraph(row[4], table_cell_style)
            st_p = Paragraph(row[5], table_ok_style)
            table_data.append([id_p, level_p, find_p, sol_p, file_p, st_p])

    col_widths = [42, 28, 140, 150, 110, 70]
    t = Table(table_data, colWidths=col_widths, repeatRows=1)
    
    t_style = TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), c_obsidian),
        ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
        ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
        ('TOPPADDING', (0, 0), (-1, -1), 3),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 3),
        ('LEFTPADDING', (0, 0), (-1, -1), 4),
        ('RIGHTPADDING', (0, 0), (-1, -1), 4),
        ('GRID', (0, 0), (-1, -1), 0.5, colors.HexColor("#CBD5E1")),
    ])

    for i in range(1, len(raw_table)):
        bg = colors.HexColor("#F8FAFC") if i % 2 == 0 else colors.white
        t_style.add('BACKGROUND', (0, i), (-1, i), bg)

    t.setStyle(t_style)
    story.append(t)

    # Dictamen Final
    story.append(Spacer(1, 10))
    dict_table_data = [[
        Paragraph(
            "<b>DICTAMEN DE AUDITORÍA: APROBADO PARA PILOTO SAAS 2026</b><br/>"
            "La plataforma Property OS cuenta con el 100% de las observaciones P0 y P1 levantadas, "
            "compilación limpia (Exit Code 0) y certificación de aislamiento multi-tenant lista para su despliegue comercial.",
            ParagraphStyle(
                'DictStyle',
                fontName='Helvetica-Bold',
                fontSize=8.5,
                leading=11.5,
                textColor=c_obsidian
            )
        )
    ]]
    dict_t = Table(dict_table_data, colWidths=[540])
    dict_t.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, -1), colors.HexColor("#FEF3C7")),
        ('BORDER', (0, 0), (-1, -1), 1.5, c_gold),
        ('TOPPADDING', (0, 0), (-1, -1), 6),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
        ('LEFTPADDING', (0, 0), (-1, -1), 8),
        ('RIGHTPADDING', (0, 0), (-1, -1), 8),
    ]))
    story.append(dict_t)

    doc.build(story)
    print(f"[OK] Documento PDF generado exitosamente: {filename}")

if __name__ == "__main__":
    generate_word_document()
    generate_pdf_document()
